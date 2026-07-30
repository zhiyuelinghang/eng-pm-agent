"""AgentScope tools backed by Dobby's session-bound internal gateway."""

import asyncio
import csv
import io
import json
import logging
import os
from typing import Any

import httpx

from agentscope.message import TextBlock, ToolResultState
from agentscope.permission import (
    PermissionBehavior,
    PermissionContext,
    PermissionDecision,
)
from agentscope.tool import ToolBase, ToolChunk


logger = logging.getLogger(__name__)

_DEFAULT_GATEWAY_URL = (
    "http://127.0.0.1:38430/api/internal/agent-tools"
)


READ_TOOL_DEFINITIONS: tuple[dict[str, Any], ...] = (
    {
        "name": "dobby_get_project_overview",
        "operation": "get_project_overview",
        "description": (
            "读取当前平台会话所绑定项目的概览、当前账号身份、进度摘要和"
            "任务/风险/资料数量。项目和账号由平台会话自动确定，不能跨项目。"
            "在回答项目现状、准备执行写操作或不确定当前权限时应优先调用。"
        ),
        "schema": {
            "type": "object",
            "properties": {},
            "additionalProperties": False,
        },
    },
    {
        "name": "dobby_list_project_items",
        "operation": "list_project_items",
        "description": (
            "读取当前项目的结构化业务数据。可查询任务、WBS 工序、风险源、"
            "质量指标、项目成员、最新信息、项目变更、通知、工程资料和日报。"
            "仅返回当前会话账号有权访问的项目数据。"
        ),
        "schema": {
            "type": "object",
            "properties": {
                "resource": {
                    "type": "string",
                    "enum": [
                        "tasks",
                        "wbs",
                        "risks",
                        "quality",
                        "members",
                        "information",
                        "changes",
                        "notifications",
                        "documents",
                        "daily_reports",
                    ],
                    "description": "要读取的业务资源类型。",
                },
                "keyword": {
                    "type": ["string", "null"],
                    "description": "可选关键词，用于名称或内容筛选。",
                    "maxLength": 200,
                },
                "status": {
                    "type": ["string", "null"],
                    "description": (
                        "可选状态筛选；查询 documents 时代表资料分类，"
                        "查询 notifications 时可用 read/unread。"
                    ),
                    "maxLength": 64,
                },
                "limit": {
                    "type": "integer",
                    "minimum": 1,
                    "maximum": 100,
                    "default": 50,
                },
            },
            "required": ["resource"],
            "additionalProperties": False,
        },
    },
    {
        "name": "dobby_search_documents",
        "operation": "search_documents",
        "description": (
            "按关键词检索当前项目已经归档的资料名称和已提取正文，返回命中"
            "文件及短摘要。此工具检索 Dobby 平台工程资料；AgentScope 知识库"
            "内容应使用 search_knowledge。"
        ),
        "schema": {
            "type": "object",
            "properties": {
                "keyword": {
                    "type": "string",
                    "minLength": 1,
                    "maxLength": 200,
                    "description": "资料检索关键词。",
                },
                "limit": {
                    "type": "integer",
                    "minimum": 1,
                    "maximum": 50,
                    "default": 20,
                },
            },
            "required": ["keyword"],
            "additionalProperties": False,
        },
    },
)


INITIALIZATION_TOOL_DEFINITIONS: tuple[dict[str, Any], ...] = (
    {
        "name": "dobby_get_project_initialization_state",
        "operation": "get_project_initialization_state",
        "description": (
            "读取当前项目已经正式入库的工程基本信息、人员、WBS、风险源、"
            "工序质量指标及最新初始化草稿摘要。此工具不返回草稿 payload；"
            "latest_draft 不为空时，必须继续调用草稿读取工具取得相关分区，"
            "不能把正式表为空解释成草稿为空。项目名称只读，不能由初始化"
            "智能体修改。"
        ),
        "schema": {
            "type": "object",
            "properties": {},
            "additionalProperties": False,
        },
    },
    {
        "name": "dobby_submit_project_initialization_draft",
        "operation": "submit_project_initialization_draft",
        "description": (
            "首次创建结构化初始化草稿，或在用户明确要求从全部原始资料重新"
            "构建时整体替换草稿。已有草稿的日常补充必须改用增量更新工具，"
            "避免覆盖其他分区。该工具只保存待核对草稿，不会修改正式项目"
            "数据；必须保留附件原值，无法确认的内容留空，"
            "不得编造。数值 0 是有效原值，必须保留为 0，不能按空值或缺失值"
            "处理。附件中每一条带 WBS 编码的记录都必须提交，即使名称疑似"
            "占位内容也不得静默丢弃。WBS 编码只表示层级和同级自然顺序，不"
            "表示前置依赖；前置关系只能来自附件或用户明确提供的信息，不得"
            "根据编码、名称或日期推断。同一上级下按 WBS 编码自然排序后，"
            "计划开始时间不得倒退；发现冲突时保留原值并交给用户核对，不得"
            "自行解释或改写。返回任何错误或警告时，必须简要说明问题，并明确"
            "提示用户点击 Dobby 平台中的“核对草稿”查看、修正或确认。"
        ),
        "schema": {
            "type": "object",
            "properties": {
                "draft_id": {
                    "type": ["integer", "null"],
                    "description": "更新已有草稿时填写；首次提交传 null。",
                },
                "source_files": {
                    "type": "array",
                    "items": {"type": "string"},
                    "maxItems": 100,
                    "description": "本次识别所依据的附件文件名。",
                },
                "payload": {
                    "type": "object",
                    "properties": {
                        "project": {
                            "type": "object",
                            "description": "项目名称不在此对象中，名称不可修改。",
                            "properties": {
                                "engineering_type_description": {"type": ["string", "null"]},
                                "contract_start_date": {"type": ["string", "null"], "format": "date"},
                                "contract_end_date": {"type": ["string", "null"], "format": "date"},
                                "contract_duration_days": {"type": ["integer", "null"], "minimum": 1},
                                "contract_amount_wan_yuan": {"type": ["number", "null"], "minimum": 0},
                                "construction_unit_name": {"type": ["string", "null"]},
                                "general_contractor_unit_name": {"type": ["string", "null"]},
                                "supervision_unit_name": {"type": ["string", "null"]},
                                "design_unit_name": {"type": ["string", "null"]},
                                "survey_unit_name": {"type": ["string", "null"]},
                            },
                            "additionalProperties": False,
                        },
                        "personnel": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "serial_no": {"type": "integer", "minimum": 1},
                                    "real_name": {"type": "string"},
                                    "identity_card_no": {"type": "string"},
                                    "position_name": {"type": "string"},
                                    "certificate_no": {"type": "string"},
                                    "responsibility_description": {"type": "string"},
                                },
                                "required": [
                                    "serial_no",
                                    "real_name",
                                    "identity_card_no",
                                    "position_name",
                                    "certificate_no",
                                    "responsibility_description",
                                ],
                                "additionalProperties": False,
                            },
                        },
                        "wbs": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "wbs_code": {"type": "string"},
                                    "parent_wbs_code": {
                                        "type": ["string", "null"],
                                        "description": (
                                            "仅按点分 WBS 编码的直接前缀确定；"
                                            "根节点传 null，不得填写“顶级 WBS”"
                                            "等说明文字。"
                                        ),
                                    },
                                    "predecessor_wbs_codes": {
                                        "type": "array",
                                        "items": {"type": "string"},
                                        "description": (
                                            "只填写附件或用户明确给出的前置 WBS"
                                            " 编码；未明确提供时传空数组，禁止按"
                                            "编码顺序、名称或日期推断。"
                                        ),
                                    },
                                    "sort_order": {"type": "integer", "minimum": 0},
                                    "color_value": {"type": ["string", "null"]},
                                    "name": {"type": "string"},
                                    "assigned_to_text": {"type": ["string", "null"]},
                                    "planned_start_at": {"type": ["string", "null"], "format": "date-time"},
                                    "planned_finish_at": {"type": ["string", "null"], "format": "date-time"},
                                    "deadline_at": {"type": ["string", "null"], "format": "date-time"},
                                    "progress_percent": {
                                        "type": ["number", "null"],
                                        "minimum": 0,
                                        "maximum": 100,
                                        "description": (
                                            "附件单元格为 0 时必须传 0；只有原始"
                                            "单元格为空时才传 null。"
                                        ),
                                    },
                                    "duration_hours": {"type": ["number", "null"], "minimum": 0},
                                    "estimated_hours": {"type": ["number", "null"], "minimum": 0},
                                    "time_log_minutes": {"type": ["integer", "null"], "minimum": 0},
                                    "status_text": {"type": ["string", "null"]},
                                    "priority_text": {"type": ["string", "null"]},
                                    "description": {"type": ["string", "null"]},
                                    "budget": {"type": ["number", "null"], "minimum": 0},
                                    "actual_cost": {"type": ["number", "null"], "minimum": 0},
                                    "msp_uid": {"type": ["string", "null"]},
                                    "msp_id": {"type": ["string", "null"]},
                                    "source_created_at": {"type": ["string", "null"], "format": "date-time"},
                                    "source_creator": {"type": ["string", "null"]},
                                    "item_type": {"type": ["string", "null"]},
                                    "source_project_path": {"type": ["string", "null"]},
                                    "level": {
                                        "type": "integer",
                                        "minimum": 1,
                                        "description": "WBS 点分编码的段数。",
                                    },
                                },
                                "required": [
                                    "wbs_code",
                                    "parent_wbs_code",
                                    "name",
                                    "planned_start_at",
                                    "planned_finish_at",
                                    "progress_percent",
                                    "status_text",
                                    "priority_text",
                                    "level",
                                ],
                                "additionalProperties": False,
                            },
                        },
                        "risks": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "serial_no": {"type": "integer", "minimum": 1},
                                    "related_wbs_code": {"type": ["string", "null"]},
                                    "related_process_name": {"type": "string"},
                                    "risk_part": {"type": "string"},
                                    "risk_level": {"type": "string"},
                                    "evaluation_condition": {"type": "string"},
                                    "risk_window_start_date": {"type": ["string", "null"], "format": "date"},
                                    "risk_window_end_date": {"type": ["string", "null"], "format": "date"},
                                    "summary": {"type": ["string", "null"]},
                                },
                                "required": [
                                    "serial_no",
                                    "related_process_name",
                                    "risk_part",
                                    "risk_level",
                                    "evaluation_condition",
                                ],
                                "additionalProperties": False,
                            },
                        },
                        "quality_requirements": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "wbs_code": {"type": "string"},
                                    "quality_acceptance_item": {"type": "string"},
                                    "control_indicator": {"type": "string"},
                                    "inspection_frequency": {"type": "string"},
                                    "related_documents": {"type": "string"},
                                },
                                "required": [
                                    "wbs_code",
                                    "quality_acceptance_item",
                                    "control_indicator",
                                    "inspection_frequency",
                                    "related_documents",
                                ],
                                "additionalProperties": False,
                            },
                        },
                    },
                    "required": [
                        "project",
                        "personnel",
                        "wbs",
                        "risks",
                        "quality_requirements",
                    ],
                    "additionalProperties": False,
                },
            },
            "required": ["payload"],
            "additionalProperties": False,
        },
    },
)


_INITIALIZATION_PAYLOAD_SCHEMA = INITIALIZATION_TOOL_DEFINITIONS[1]["schema"][
    "properties"
]["payload"]

INITIALIZATION_TOOL_DEFINITIONS += (
    {
        "name": "dobby_get_project_initialization_draft",
        "operation": "get_project_initialization_draft",
        "description": (
            "读取当前初始化会话最新草稿的完整结构化数据。必须按分区读取；"
            "WBS、人员、风险、质量和校验结果支持 start/limit 分页。补充已有"
            "草稿前，先读取本次所需分区；不得因为正式业务表尚未入库就重新"
            "解析全部旧附件。草稿中的 WBS 可以直接用于匹配新增风险和质量"
            "指标。"
        ),
        "schema": {
            "type": "object",
            "properties": {
                "section": {
                    "type": "string",
                    "enum": [
                        "project",
                        "personnel",
                        "wbs",
                        "risks",
                        "quality_requirements",
                        "validation_issues",
                    ],
                    "description": "需要读取的草稿分区。",
                },
                "start": {
                    "type": "integer",
                    "minimum": 1,
                    "default": 1,
                    "description": "数组分区从 1 开始的起始记录。",
                },
                "limit": {
                    "type": "integer",
                    "minimum": 1,
                    "maximum": 500,
                    "default": 100,
                    "description": "本次最多返回的记录数。",
                },
            },
            "required": ["section"],
            "additionalProperties": False,
        },
    },
    {
        "name": "dobby_update_project_initialization_draft",
        "operation": "update_project_initialization_draft",
        "description": (
            "增量更新已有初始化草稿。只替换 patch 中明确提供的分区，未提供"
            "的工程信息、人员、WBS、风险和质量分区保持不变；project 对象按"
            "字段合并，人员、WBS、风险和质量等列表分区整体替换。更新前必须"
            "读取草稿并传入最新 expected_revision，防止覆盖并发修改。列表"
            "分区传空数组表示明确清空；不得传 null。新增风险或质量指标时，"
            "应先读取草稿 WBS 完成编码匹配，无需重新读取旧 WBS 附件。返回"
            "任何错误或警告时，必须提示用户点击“核对草稿”。"
        ),
        "schema": {
            "type": "object",
            "properties": {
                "draft_id": {
                    "type": "integer",
                    "minimum": 1,
                    "description": "读取草稿工具返回的草稿 ID。",
                },
                "expected_revision": {
                    "type": "integer",
                    "minimum": 1,
                    "description": "读取草稿工具返回的当前修订号。",
                },
                "source_files": {
                    "type": "array",
                    "items": {"type": "string"},
                    "maxItems": 100,
                    "description": "本次增量识别新增使用的附件文件名。",
                },
                "patch": {
                    "type": "object",
                    "properties": _INITIALIZATION_PAYLOAD_SCHEMA["properties"],
                    "minProperties": 1,
                    "additionalProperties": False,
                    "description": (
                        "只填写本次需要替换的草稿分区；未填写的分区保持不变。"
                    ),
                },
            },
            "required": ["draft_id", "expected_revision", "patch"],
            "additionalProperties": False,
        },
    },
)


INITIALIZATION_FILE_TOOL_DEFINITION: dict[str, Any] = {
    "name": "dobby_read_project_initialization_file",
    "description": (
        "读取并解析当前项目初始化会话中已上传的原始附件。平台只负责授权"
        "和传输原始文件，XLSX、DOCX、PDF、CSV、TXT/Markdown 的解析在"
        "AgentScope 进程内完成。先用初始化状态工具取得 file_id；大文件应"
        "按 start 和 limit 分段读取，XLSX 可通过 sheet_name 指定工作表。"
    ),
    "schema": {
        "type": "object",
        "properties": {
            "file_id": {
                "type": "integer",
                "description": "初始化状态中返回的附件 ID。",
            },
            "sheet_name": {
                "type": ["string", "null"],
                "description": "读取 XLSX 时指定工作表；不填则读取第一张表。",
            },
            "start": {
                "type": "integer",
                "minimum": 1,
                "default": 1,
                "description": "起始行、页或文档块，均从 1 开始。",
            },
            "limit": {
                "type": "integer",
                "minimum": 1,
                "maximum": 500,
                "default": 100,
                "description": "本次最多返回的行、页或文档块数量。",
            },
        },
        "required": ["file_id"],
        "additionalProperties": False,
    },
}


WRITE_TOOL_DEFINITIONS: tuple[dict[str, Any], ...] = (
    {
        "name": "dobby_create_task",
        "operation": "create_task",
        "description": (
            "在当前项目创建一条真实任务，并写入审计记录。调用前应先查询成员、"
            "WBS 和风险源，引用的 ID 必须属于当前项目。不要把“建议创建”当作"
            "已经创建；只有工具成功返回后才能向用户确认完成。"
        ),
        "schema": {
            "type": "object",
            "properties": {
                "title": {"type": "string", "minLength": 1, "maxLength": 300},
                "task_type": {
                    "type": "string",
                    "enum": [
                        "risk_alert",
                        "material_missing",
                        "daily_confirm",
                        "draft_review",
                        "fill_platform",
                    ],
                    "default": "risk_alert",
                },
                "risk_level": {
                    "type": "string",
                    "enum": ["low", "medium", "high", "critical"],
                    "default": "low",
                },
                "assignee_user_id": {"type": ["integer", "null"]},
                "confirmer_user_id": {"type": ["integer", "null"]},
                "due_at": {
                    "type": ["string", "null"],
                    "description": "截止日期或时间，推荐 ISO 格式。",
                },
                "wbs_item_id": {"type": ["integer", "null"]},
                "risk_source_id": {"type": ["integer", "null"]},
                "trigger_reason": {
                    "type": ["string", "null"],
                    "maxLength": 2000,
                },
                "required_materials": {
                    "type": "array",
                    "items": {"type": "string"},
                    "maxItems": 30,
                    "default": [],
                },
            },
            "required": ["title"],
            "additionalProperties": False,
        },
    },
    {
        "name": "dobby_update_task",
        "operation": "update_task",
        "description": (
            "更新当前项目的真实任务：状态流转、转交、添加处置说明或更新流程"
            "步骤。必须先读取任务当前状态；工具会执行合法流转和项目成员校验。"
        ),
        "schema": {
            "type": "object",
            "properties": {
                "task_id": {"type": "integer"},
                "action": {
                    "type": "string",
                    "enum": [
                        "transition",
                        "reassign",
                        "add_note",
                        "update_step",
                    ],
                },
                "status": {
                    "type": ["string", "null"],
                    "description": "状态流转或步骤更新时的新状态。",
                },
                "assignee_user_id": {"type": ["integer", "null"]},
                "step_index": {
                    "type": ["integer", "null"],
                    "minimum": 0,
                    "description": "更新流程步骤时使用，从 0 开始。",
                },
                "note": {"type": ["string", "null"], "maxLength": 2000},
            },
            "required": ["task_id", "action"],
            "additionalProperties": False,
        },
    },
    {
        "name": "dobby_dispose_information",
        "operation": "dispose_information",
        "description": (
            "确认、否认或修订当前项目的一条最新信息记录，并写入审计日志。"
            "修订时必须提供 content。执行前先读取原记录和可用依据。"
        ),
        "schema": {
            "type": "object",
            "properties": {
                "record_id": {"type": "integer"},
                "action": {
                    "type": "string",
                    "enum": ["confirm", "deny", "revise"],
                },
                "content": {"type": ["string", "null"], "maxLength": 10000},
            },
            "required": ["record_id", "action"],
            "additionalProperties": False,
        },
    },
    {
        "name": "dobby_create_project_change",
        "operation": "create_project_change",
        "description": (
            "在当前项目登记一条待审核的项目变更，并保留来源引用和审计记录。"
            "仅用于已有事实或用户明确要求登记的变更，不得自行编造变更事实。"
        ),
        "schema": {
            "type": "object",
            "properties": {
                "category": {
                    "type": "string",
                    "minLength": 1,
                    "maxLength": 100,
                    "default": "工程内容变更",
                },
                "title": {"type": "string", "minLength": 1, "maxLength": 300},
                "content": {
                    "type": "string",
                    "minLength": 1,
                    "maxLength": 10000,
                },
                "source_refs": {
                    "type": "array",
                    "items": {"type": "string"},
                    "maxItems": 30,
                    "default": [],
                },
            },
            "required": ["title", "content"],
            "additionalProperties": False,
        },
    },
    {
        "name": "dobby_update_document_category",
        "operation": "update_document_category",
        "description": (
            "修改当前项目一份已归档资料的分类，并写入审计记录。先查询资料 ID，"
            "仅在用户要求或分类依据充分时执行。"
        ),
        "schema": {
            "type": "object",
            "properties": {
                "attachment_id": {"type": "integer"},
                "category": {
                    "type": "string",
                    "minLength": 1,
                    "maxLength": 100,
                },
            },
            "required": ["attachment_id", "category"],
            "additionalProperties": False,
        },
    },
)


ADMIN_WRITE_TOOL_DEFINITIONS: tuple[dict[str, Any], ...] = (
    {
        "name": "dobby_create_risk",
        "operation": "create_risk",
        "description": (
            "以当前平台管理员身份在当前项目新增真实风险源，并写入审计记录。"
            "这是管理员级写操作；必须使用风险清单的新字段并先核验关联 WBS。"
        ),
        "schema": {
            "type": "object",
            "properties": {
                "serial_no": {"type": "integer", "minimum": 1},
                "related_wbs_item_id": {"type": ["integer", "null"]},
                "related_process_name": {"type": "string", "minLength": 1, "maxLength": 300},
                "risk_part": {"type": "string", "minLength": 1, "maxLength": 300},
                "risk_level": {"type": "string", "minLength": 1, "maxLength": 50},
                "evaluation_condition": {"type": "string", "minLength": 1, "maxLength": 20000},
                "risk_window_start_date": {"type": ["string", "null"], "format": "date"},
                "risk_window_end_date": {"type": ["string", "null"], "format": "date"},
                "summary": {"type": ["string", "null"], "maxLength": 20000},
            },
            "required": [
                "serial_no",
                "related_process_name",
                "risk_part",
                "risk_level",
                "evaluation_condition",
            ],
            "additionalProperties": False,
        },
    },
    {
        "name": "dobby_update_wbs_progress",
        "operation": "update_wbs_progress",
        "description": (
            "以当前平台管理员身份更新当前项目 WBS 工序的真实进度和状态，"
            "并写入审计记录。必须先读取工序现状和进度依据。"
        ),
        "schema": {
            "type": "object",
            "properties": {
                "wbs_item_id": {"type": "integer"},
                "progress_percent": {
                    "type": "number",
                    "minimum": 0,
                    "maximum": 100,
                },
                "status_text": {"type": ["string", "null"], "maxLength": 100},
                "note": {"type": ["string", "null"], "maxLength": 1000},
            },
            "required": ["wbs_item_id", "progress_percent"],
            "additionalProperties": False,
        },
    },
)


def _json_cell(value: Any) -> Any:
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return str(value)


def _decode_text_file(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _parse_initialization_file(
    content: bytes,
    suffix: str,
    sheet_name: str | None,
    start: int,
    limit: int,
) -> dict[str, Any]:
    """Parse a bounded portion of an initialization file in AgentScope."""
    if suffix == ".xlsx":
        from openpyxl import load_workbook

        workbook = load_workbook(
            io.BytesIO(content),
            read_only=True,
            data_only=True,
        )
        sheets = list(workbook.sheetnames)
        selected_sheet = sheet_name or (sheets[0] if sheets else None)
        if selected_sheet is None or selected_sheet not in sheets:
            workbook.close()
            raise ValueError(
                f"工作表不存在，可选工作表：{json.dumps(sheets, ensure_ascii=False)}",
            )
        worksheet = workbook[selected_sheet]
        end = min(worksheet.max_row, start + limit - 1)
        rows = [
            [_json_cell(cell) for cell in row]
            for row in worksheet.iter_rows(
                min_row=start,
                max_row=end,
                values_only=True,
            )
        ]
        result = {
            "format": "xlsx",
            "sheet_names": sheets,
            "sheet_name": selected_sheet,
            "start_row": start,
            "end_row": end,
            "total_rows": worksheet.max_row,
            "rows": rows,
            "next_start": end + 1 if end < worksheet.max_row else None,
        }
        workbook.close()
        return result

    if suffix == ".csv":
        rows = list(csv.reader(io.StringIO(_decode_text_file(content))))
        selected = rows[start - 1:start - 1 + limit]
        end = start + len(selected) - 1
        return {
            "format": "csv",
            "start_row": start,
            "end_row": end,
            "total_rows": len(rows),
            "rows": selected,
            "next_start": end + 1 if end < len(rows) else None,
        }

    if suffix in {".txt", ".md"}:
        lines = _decode_text_file(content).splitlines()
        selected = lines[start - 1:start - 1 + limit]
        end = start + len(selected) - 1
        return {
            "format": suffix.removeprefix("."),
            "start_line": start,
            "end_line": end,
            "total_lines": len(lines),
            "lines": selected,
            "next_start": end + 1 if end < len(lines) else None,
        }

    if suffix == ".docx":
        from docx import Document

        document = Document(io.BytesIO(content))
        blocks: list[dict[str, Any]] = [
            {"type": "paragraph", "text": paragraph.text}
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                blocks.append(
                    {
                        "type": "table_row",
                        "table": table_index,
                        "row": row_index,
                        "cells": [cell.text for cell in row.cells],
                    },
                )
        selected = blocks[start - 1:start - 1 + limit]
        end = start + len(selected) - 1
        return {
            "format": "docx",
            "start_block": start,
            "end_block": end,
            "total_blocks": len(blocks),
            "blocks": selected,
            "next_start": end + 1 if end < len(blocks) else None,
        }

    if suffix == ".pdf":
        import pdfplumber

        with pdfplumber.open(io.BytesIO(content)) as document:
            total_pages = len(document.pages)
            page_limit = min(limit, 20)
            selected_pages = document.pages[start - 1:start - 1 + page_limit]
            pages = [
                {
                    "page": start + index,
                    "text": (page.extract_text() or "")[:50000],
                }
                for index, page in enumerate(selected_pages)
            ]
        end = start + len(pages) - 1
        return {
            "format": "pdf",
            "start_page": start,
            "end_page": end,
            "total_pages": total_pages,
            "pages": pages,
            "next_start": end + 1 if end < total_pages else None,
        }

    raise ValueError(f"不支持的初始化附件格式：{suffix or '未知'}")


class DobbyInitializationFileTool(ToolBase):
    """AgentScope-side parser for raw, session-authorized platform files."""

    name = str(INITIALIZATION_FILE_TOOL_DEFINITION["name"])
    description = str(INITIALIZATION_FILE_TOOL_DEFINITION["description"])
    input_schema = dict(INITIALIZATION_FILE_TOOL_DEFINITION["schema"])
    is_concurrency_safe = False
    is_external_tool = False
    is_state_injected = False
    is_mcp = False
    is_read_only = True

    def __init__(
        self,
        session_id: str,
        base_url: str,
        token: str,
    ) -> None:
        super().__init__()
        self._session_id = session_id
        self._base_url = base_url.rstrip("/")
        self._token = token

    async def check_permissions(
        self,
        tool_input: dict[str, Any],
        context: PermissionContext,
    ) -> PermissionDecision:
        del tool_input, context
        return PermissionDecision(
            behavior=PermissionBehavior.ALLOW,
            message="该工具只读取当前初始化会话已授权的原始附件。",
        )

    async def call(
        self,
        file_id: int,
        sheet_name: str | None = None,
        start: int = 1,
        limit: int = 100,
    ) -> ToolChunk:
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.get(
                    (
                        f"{self._base_url}/initialization-files/"
                        f"{file_id}/content"
                    ),
                    headers={"Authorization": f"Bearer {self._token}"},
                    params={
                        "agentscope_session_id": self._session_id,
                    },
                )
            if response.is_error:
                try:
                    detail = response.json().get("detail", response.text)
                except ValueError:
                    detail = response.text or response.reason_phrase
                raise ValueError(
                    f"读取附件失败（{response.status_code}）：{detail}",
                )
            suffix = response.headers.get(
                "X-Dobby-File-Extension",
                "",
            ).lower()
            parsed = await asyncio.to_thread(
                _parse_initialization_file,
                response.content,
                suffix,
                sheet_name,
                max(1, start),
                min(max(1, limit), 500),
            )
            parsed["file_id"] = file_id
            return ToolChunk(
                content=[
                    TextBlock(
                        text=json.dumps(
                            parsed,
                            ensure_ascii=False,
                            default=str,
                        ),
                    ),
                ],
                state=ToolResultState.SUCCESS,
                is_last=True,
                metadata={
                    "operation": "read_project_initialization_file",
                    "platform_data_changed": False,
                },
            )
        except (
            httpx.HTTPError,
            ImportError,
            ValueError,
            OSError,
            RuntimeError,
        ) as exc:
            logger.warning("Unable to parse initialization file: %s", exc)
            return ToolChunk(
                content=[TextBlock(text=f"初始化附件读取失败：{exc}")],
                state=ToolResultState.ERROR,
                is_last=True,
                metadata={
                    "operation": "read_project_initialization_file",
                },
            )


class DobbyGatewayTool(ToolBase):
    """One semantic Dobby operation bound to a platform chat session."""

    is_concurrency_safe = False
    is_external_tool = False
    is_state_injected = False
    is_mcp = False

    def __init__(
        self,
        definition: dict[str, Any],
        session_id: str,
        base_url: str,
        token: str,
        requires_confirmation: bool,
        changes_business_data: bool,
    ) -> None:
        super().__init__()
        self.name = str(definition["name"])
        self.description = str(definition["description"])
        self.input_schema = dict(definition["schema"])
        self.is_read_only = not changes_business_data
        self._operation = str(definition["operation"])
        self._session_id = session_id
        self._base_url = base_url.rstrip("/")
        self._token = token
        self._requires_confirmation = requires_confirmation
        self._changes_business_data = changes_business_data

    async def check_permissions(
        self,
        tool_input: dict[str, Any],
        context: PermissionContext,
    ) -> PermissionDecision:
        del tool_input, context
        if not self._requires_confirmation:
            return PermissionDecision(
                behavior=PermissionBehavior.ALLOW,
                message="该 Dobby 工具仅读取项目数据或保存待核对草稿。",
            )
        return PermissionDecision(
            behavior=PermissionBehavior.ASK,
            message=(
                "该操作将修改当前平台账号有权访问的项目数据，"
                "需要权限模式或权限审核员批准。"
            ),
        )

    async def call(self, **kwargs: Any) -> ToolChunk:
        try:
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.post(
                    f"{self._base_url}/execute",
                    headers={"Authorization": f"Bearer {self._token}"},
                    json={
                        "agentscope_session_id": self._session_id,
                        "operation": self._operation,
                        "arguments": kwargs,
                    },
                )
            if response.is_error:
                try:
                    payload = response.json()
                    detail = payload.get("detail", payload)
                except ValueError:
                    detail = response.text or response.reason_phrase
                return ToolChunk(
                    content=[
                        TextBlock(
                            text=(
                                f"Dobby 操作失败（{response.status_code}）："
                                f"{json.dumps(detail, ensure_ascii=False)}"
                            ),
                        ),
                    ],
                    state=ToolResultState.ERROR,
                    is_last=True,
                    metadata={"operation": self._operation},
                )
            payload = response.json()
            return ToolChunk(
                content=[
                    TextBlock(
                        text=json.dumps(payload, ensure_ascii=False, default=str),
                    ),
                ],
                state=ToolResultState.SUCCESS,
                is_last=True,
                metadata={
                    "operation": self._operation,
                    "platform_data_changed": self._changes_business_data,
                },
            )
        except httpx.HTTPError as exc:
            logger.warning("Dobby agent tool gateway request failed: %s", exc)
            return ToolChunk(
                content=[
                    TextBlock(
                        text=f"Dobby 平台工具网关当前不可用：{exc}",
                    ),
                ],
                state=ToolResultState.ERROR,
                is_last=True,
                metadata={"operation": self._operation},
            )


def _gateway_settings() -> tuple[str, str]:
    base_url = os.getenv("DOBBY_AGENT_TOOL_BASE_URL", _DEFAULT_GATEWAY_URL).strip()
    token = (
        os.getenv("DOBBY_AGENT_TOOL_TOKEN", "").strip()
        or os.getenv("AGENTSCOPE_SERVICE_TOKEN", "").strip()
    )
    return base_url, token


async def create_dobby_agent_tools(
    user_id: str,
    agent_id: str,
    session_id: str,
) -> list[ToolBase]:
    """Build tools only for sessions created through the Dobby platform."""
    del user_id
    base_url, token = _gateway_settings()
    if not token:
        logger.warning(
            "Dobby agent tools disabled: no internal service token configured.",
        )
        return []

    try:
        async with httpx.AsyncClient(timeout=3.0) as client:
            response = await client.get(
                f"{base_url}/context",
                headers={"Authorization": f"Bearer {token}"},
                params={"agentscope_session_id": session_id},
            )
        if response.status_code in {403, 404}:
            return []
        response.raise_for_status()
        context = response.json().get("data") or {}
    except (httpx.HTTPError, ValueError, TypeError) as exc:
        logger.warning("Unable to resolve Dobby tool context: %s", exc)
        return []

    # Prevent a mapped session from being used to assemble tools for a
    # different AgentScope agent.
    if str(context.get("agent_id") or "") != agent_id:
        return []

    capabilities = context.get("capabilities") or {}
    definitions: list[tuple[dict[str, Any], bool, bool]] = [
        *((definition, False, False) for definition in READ_TOOL_DEFINITIONS),
    ]
    if capabilities.get("initialization_draft"):
        definitions.extend(
            (definition, False, False)
            for definition in INITIALIZATION_TOOL_DEFINITIONS
        )
    if capabilities.get("write"):
        definitions.extend(
            (definition, True, True) for definition in WRITE_TOOL_DEFINITIONS
        )
    if capabilities.get("admin_write"):
        definitions.extend(
            (definition, True, True)
            for definition in ADMIN_WRITE_TOOL_DEFINITIONS
        )

    tools: list[ToolBase] = [
        DobbyGatewayTool(
            definition=definition,
            session_id=session_id,
            base_url=base_url,
            token=token,
            requires_confirmation=requires_confirmation,
            changes_business_data=changes_business_data,
        )
        for (
            definition,
            requires_confirmation,
            changes_business_data,
        ) in definitions
    ]
    if capabilities.get("initialization_draft"):
        tools.append(
            DobbyInitializationFileTool(
                session_id=session_id,
                base_url=base_url,
                token=token,
            ),
        )
    return tools
